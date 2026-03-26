#!/usr/bin/env python3
"""
WOCS 견적서 자동 생성기
- 천막구조물 / 글램핑 두 가지 견적서 지원
- Word(.docx) + PDF 동시 생성
- 터미널 입력 → 자동 생성
"""

import os
import sys
import json
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers

# ============================================
# 회사 정보
# ============================================
COMPANY = {
    "name": "WOCS (우성어닝천막공사캠프시스템)",
    "ceo": "김우성",
    "phone": "010-4337-0582",
    "email": "info@wocs.kr",
    "address": "전남 화순군 사평면 유마로 592",
    "website": "wocs.kr",
    "biz_no": "413-30-01052",
}

# ============================================
# 단가표 (내장)
# ============================================
PRICES = {
    "천막구조물": {
        "어닝": {"unit": "㎡", "base_price": 150000},
        "천막": {"unit": "㎡", "base_price": 120000},
        "파고라": {"unit": "㎡", "base_price": 200000},
        "방풍막": {"unit": "㎡", "base_price": 100000},
        "렉산": {"unit": "㎡", "base_price": 180000},
        "철구조물": {"unit": "㎡", "base_price": 250000},
    },
    "글램핑": {
        "사파리텐트": {"unit": "동", "base_price": 8500000},
        "돔텐트": {"unit": "동", "base_price": 6500000},
        "몽골텐트": {"unit": "동", "base_price": 5500000},
        "시그니처텐트": {"unit": "동", "base_price": 12000000},
    },
    "글램핑옵션": {
        "기초공사": {"unit": "동", "base_price": 1500000},
        "데크": {"unit": "동", "base_price": 2000000},
        "전기공사": {"unit": "동", "base_price": 800000},
        "수도공사": {"unit": "동", "base_price": 600000},
        "조인트특허시스템": {"unit": "동", "base_price": 500000},
    },
}

FRAME_MATERIALS = ["스틸", "알루미늄", "스테인리스", "아연도금"]
FABRIC_TYPES = ["PVC 코팅", "방염포", "타포린", "투명 PVC", "어닝 원단"]
COLORS = ["화이트", "아이보리", "그레이", "블랙", "그린", "베이지", "커스텀"]

# ============================================
# 출력 디렉토리
# ============================================
OUTPUT_DIR = Path(__file__).parent / "quotes"
LOGO_PATH = Path(__file__).parent / "assets" / "images" / "wocs-logo.png"
COUNTER_FILE = Path(__file__).parent / "quotes" / ".quote_counter.json"


def get_next_quote_number():
    """견적번호 자동 생성: WOCS-2026-001"""
    year = datetime.date.today().year
    OUTPUT_DIR.mkdir(exist_ok=True)

    counter = {}
    if COUNTER_FILE.exists():
        counter = json.loads(COUNTER_FILE.read_text(encoding="utf-8"))

    year_key = str(year)
    seq = counter.get(year_key, 0) + 1
    counter[year_key] = seq
    COUNTER_FILE.write_text(json.dumps(counter, ensure_ascii=False), encoding="utf-8")

    return f"WOCS-{year}-{seq:03d}"


def fmt_money(amount):
    """금액 포맷: 1,234,567원"""
    return f"{int(amount):,}원"


def input_default(prompt, default=""):
    """기본값 있는 입력"""
    if default:
        val = input(f"  {prompt} [{default}]: ").strip()
        return val if val else default
    return input(f"  {prompt}: ").strip()


def input_choice(prompt, choices):
    """선택지 입력"""
    print(f"\n  {prompt}")
    for i, c in enumerate(choices, 1):
        print(f"    {i}. {c}")
    while True:
        sel = input(f"  선택 (1-{len(choices)}): ").strip()
        if sel.isdigit() and 1 <= int(sel) <= len(choices):
            return choices[int(sel) - 1]
        print("  → 올바른 번호를 입력하세요")


def input_yn(prompt, default="y"):
    """Y/N 입력"""
    val = input(f"  {prompt} (y/n) [{default}]: ").strip().lower()
    return (val or default) == "y"


# ============================================
# 스타일 설정
# ============================================
def compact_paragraph(p, before=0, after=0):
    """단락 간격 축소"""
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def set_cell_style(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                   font_size=9, bg_color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    compact_paragraph(p, 1, 1)
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.bold = bold
    if bg_color:
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): bg_color,
            qn("w:val"): "clear",
        })
        shading.append(shd)


def add_header(doc, quote_number, quote_type):
    """견적서 헤더"""
    # 로고 또는 텍스트 헤더
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(2))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact_paragraph(p, 0, 0)
        run = p.add_run("WOCS")
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.bold = True
        run.font.name = "맑은 고딕"
        run = p.add_run("\n우성어닝천막공사캠프시스템")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = "맑은 고딕"

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact_paragraph(p, 2, 2)
    run = p.add_run(f"  {quote_type} 견적서  ")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = "맑은 고딕"

    # 견적 정보 라인
    today = datetime.date.today()
    expire = today + datetime.timedelta(days=30)
    info_text = (
        f"견적번호: {quote_number}   |   "
        f"작성일: {today.strftime('%Y-%m-%d')}   |   "
        f"유효기간: {expire.strftime('%Y-%m-%d')}까지"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact_paragraph(p, 0, 4)
    run = p.add_run(info_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = "맑은 고딕"


def add_customer_info(doc, info):
    """고객 정보 테이블"""
    p = doc.add_paragraph()
    compact_paragraph(p, 4, 2)
    run = p.add_run("■ 고객 정보")
    run.font.size = Pt(11)
    run.bold = True
    run.font.name = "맑은 고딕"

    table = doc.add_table(rows=len(info), cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(info):
        set_cell_style(table.cell(i, 0), key, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F0F4FF")
        set_cell_style(table.cell(i, 1), val,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
        table.cell(i, 0).width = Cm(4)
        table.cell(i, 1).width = Cm(12)


def add_item_table(doc, headers, rows, title="■ 견적 내역"):
    """견적 항목 테이블"""
    p = doc.add_paragraph()
    compact_paragraph(p, 6, 2)
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.bold = True
    run.font.name = "맑은 고딕"

    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for j, h in enumerate(headers):
        set_cell_style(table.cell(0, j), h, bold=True, bg_color="1A56DB")
        # 헤더 글자 흰색
        for run in table.cell(0, j).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 데이터
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if isinstance(val, (int, float)) or "원" in str(val) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_style(table.cell(i + 1, j), str(val), align=align)

    return table


def add_summary(doc, subtotal, install_fee, transport_fee, demolish_fee):
    """합계 요약"""
    extra = install_fee + transport_fee + demolish_fee
    supply = subtotal + extra
    vat = int(supply * 0.1)
    total = supply + vat

    p = doc.add_paragraph()
    compact_paragraph(p, 6, 2)
    run = p.add_run("■ 합계")
    run.font.size = Pt(11)
    run.bold = True
    run.font.name = "맑은 고딕"

    summary_items = [("공급가액 (품목 소계)", fmt_money(subtotal))]
    if install_fee:
        summary_items.append(("시공비", fmt_money(install_fee)))
    if transport_fee:
        summary_items.append(("운반비", fmt_money(transport_fee)))
    if demolish_fee:
        summary_items.append(("철거비", fmt_money(demolish_fee)))
    summary_items.append(("부가세 (10%)", fmt_money(vat)))
    summary_items.append(("총 합계", fmt_money(total)))

    table = doc.add_table(rows=len(summary_items), cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(summary_items):
        is_total = (i == len(summary_items) - 1)
        set_cell_style(table.cell(i, 0), label, bold=is_total,
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       bg_color="FFF3E0" if is_total else "F0F4FF")
        set_cell_style(table.cell(i, 1), val, bold=is_total,
                       align=WD_ALIGN_PARAGRAPH.RIGHT,
                       font_size=11 if is_total else 9,
                       bg_color="FFF3E0" if is_total else None)
        table.cell(i, 0).width = Cm(6)
        table.cell(i, 1).width = Cm(10)

    return total


def add_footer(doc):
    """회사 정보 푸터"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact_paragraph(p, 10, 2)
    lines = [
        f"{COMPANY['name']}",
        f"대표: {COMPANY['ceo']}  |  사업자번호: {COMPANY['biz_no']}",
        f"주소: {COMPANY['address']}",
        f"전화: {COMPANY['phone']}  |  이메일: {COMPANY['email']}  |  {COMPANY['website']}",
    ]
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = "맑은 고딕"
        if i < len(lines) - 1:
            p.add_run("\n")

    # 날인
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    compact_paragraph(p, 8, 0)
    run = p.add_run("위 금액을 견적합니다.\n")
    run.font.size = Pt(9)
    run.font.name = "맑은 고딕"
    run = p.add_run(f"{COMPANY['name']}  대표  {COMPANY['ceo']}  (인)")
    run.font.size = Pt(10)
    run.bold = True
    run.font.name = "맑은 고딕"


def save_excel(quote_number, quote_type, customer_info, items, summary_data):
    """Excel(.xlsx) 견적서 생성"""
    wb = Workbook()
    ws = wb.active
    ws.title = "견적서"

    # 스타일 정의
    header_font = Font(name="맑은 고딕", size=20, bold=True, color="1A56DB")
    sub_font = Font(name="맑은 고딕", size=9, color="555555")
    info_font = Font(name="맑은 고딕", size=8, color="888888")
    section_font = Font(name="맑은 고딕", size=11, bold=True)
    label_font = Font(name="맑은 고딕", size=9, bold=True)
    data_font = Font(name="맑은 고딕", size=9)
    total_font = Font(name="맑은 고딕", size=11, bold=True)
    stamp_font = Font(name="맑은 고딕", size=10, bold=True)

    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    blue_fill = PatternFill(start_color="1A56DB", end_color="1A56DB", fill_type="solid")
    light_fill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
    total_fill = PatternFill(start_color="FFF3E0", end_color="FFF3E0", fill_type="solid")
    white_font = Font(name="맑은 고딕", size=9, bold=True, color="FFFFFF")
    center = Alignment(horizontal="center", vertical="center")
    left = Alignment(horizontal="left", vertical="center")
    right = Alignment(horizontal="right", vertical="center")

    # 열 너비
    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 22
    ws.column_dimensions["E"].width = 8
    ws.column_dimensions["F"].width = 16
    ws.column_dimensions["G"].width = 16

    row = 1

    # 헤더
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    c = ws.cell(row=row, column=1, value="WOCS")
    c.font = header_font
    c.alignment = center
    row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    c = ws.cell(row=row, column=1, value="우성어닝천막공사캠프시스템")
    c.font = sub_font
    c.alignment = center
    row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    c = ws.cell(row=row, column=1, value=f"{quote_type} 견적서")
    c.font = Font(name="맑은 고딕", size=16, bold=True)
    c.alignment = center
    row += 1

    today = datetime.date.today()
    expire = today + datetime.timedelta(days=30)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    c = ws.cell(row=row, column=1,
                value=f"견적번호: {quote_number}  |  작성일: {today}  |  유효기간: {expire}까지")
    c.font = info_font
    c.alignment = center
    row += 2

    # 고객 정보
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    ws.cell(row=row, column=1, value="■ 고객 정보").font = section_font
    row += 1
    for key, val in customer_info:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
        c = ws.cell(row=row, column=1, value=key)
        c.font = label_font
        c.alignment = center
        c.fill = light_fill
        c.border = thin_border
        ws.cell(row=row, column=2).border = thin_border
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=7)
        c = ws.cell(row=row, column=3, value=val)
        c.font = data_font
        c.alignment = left
        for col in range(1, 8):
            ws.cell(row=row, column=col).border = thin_border
        row += 1
    row += 1

    # 견적 내역
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    ws.cell(row=row, column=1, value="■ 견적 내역").font = section_font
    row += 1

    # 테이블 헤더
    col_headers = items["headers"]
    for j, h in enumerate(col_headers, 1):
        c = ws.cell(row=row, column=j, value=h)
        c.font = white_font
        c.fill = blue_fill
        c.alignment = center
        c.border = thin_border
    row += 1

    # 데이터
    for item_row in items["rows"]:
        for j, val in enumerate(item_row, 1):
            c = ws.cell(row=row, column=j, value=val)
            c.font = data_font
            c.alignment = right if j >= 5 else center
            c.border = thin_border
        row += 1
    row += 1

    # 합계
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    ws.cell(row=row, column=1, value="■ 합계").font = section_font
    row += 1

    for i, (label, val) in enumerate(summary_data):
        is_total = (i == len(summary_data) - 1)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        c = ws.cell(row=row, column=1, value=label)
        c.font = total_font if is_total else label_font
        c.alignment = center
        c.fill = total_fill if is_total else light_fill
        c.border = thin_border
        for col in range(1, 4):
            ws.cell(row=row, column=col).border = thin_border
        ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=7)
        c = ws.cell(row=row, column=4, value=val)
        c.font = total_font if is_total else data_font
        c.alignment = right
        c.fill = total_fill if is_total else PatternFill()
        for col in range(4, 8):
            ws.cell(row=row, column=col).border = thin_border
        row += 1
    row += 1

    # 푸터
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    c = ws.cell(row=row, column=1,
                value=f"{COMPANY['name']}  |  대표: {COMPANY['ceo']}  |  사업자번호: {COMPANY['biz_no']}")
    c.font = info_font
    c.alignment = center
    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    c = ws.cell(row=row, column=1,
                value=f"주소: {COMPANY['address']}  |  전화: {COMPANY['phone']}  |  {COMPANY['email']}")
    c.font = info_font
    c.alignment = center
    row += 2

    ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=7)
    c = ws.cell(row=row, column=4, value="위 금액을 견적합니다.")
    c.font = data_font
    c.alignment = right
    row += 1
    ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=7)
    c = ws.cell(row=row, column=4,
                value=f"{COMPANY['name']}  대표  {COMPANY['ceo']}  (인)")
    c.font = stamp_font
    c.alignment = right

    # 인쇄 설정
    ws.print_area = f"A1:G{row}"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1

    OUTPUT_DIR.mkdir(exist_ok=True)
    safe_type = quote_type.replace(" ", "_")
    xlsx_path = OUTPUT_DIR / f"{quote_number}_{safe_type}.xlsx"
    wb.save(str(xlsx_path))
    print(f"  ✓ Excel 저장: {xlsx_path}")
    return xlsx_path


def save_document(doc, quote_number, quote_type):
    """docx 저장 + PDF 변환"""
    OUTPUT_DIR.mkdir(exist_ok=True)
    safe_type = quote_type.replace(" ", "_")
    base = f"{quote_number}_{safe_type}"
    docx_path = OUTPUT_DIR / f"{base}.docx"
    pdf_path = OUTPUT_DIR / f"{base}.pdf"

    doc.save(str(docx_path))
    print(f"\n  ✓ Word 저장: {docx_path}")

    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"  ✓ PDF  저장: {pdf_path}")
    except Exception as e:
        print(f"  ⚠ PDF 변환 실패 (MS Word 필요): {e}")
        print(f"    → Word 파일을 열어 직접 PDF로 저장하세요")

    return docx_path, pdf_path


# ============================================
# 천막구조물 견적서
# ============================================
def create_tent_quote():
    print("\n" + "=" * 50)
    print("  🏗️  천막구조물 견적서 작성")
    print("=" * 50)

    # 고객 정보
    print("\n[ 고객 정보 ]")
    customer = input_default("고객명")
    address = input_default("현장주소")
    schedule = input_default("시공일정", "협의 후 결정")

    # 품목 입력
    items = []
    print("\n[ 품목 입력 ] (완료시 빈 칸 Enter)")
    idx = 1
    while True:
        print(f"\n  --- 품목 {idx} ---")
        struct_types = list(PRICES["천막구조물"].keys())
        stype = input_choice("구조물 종류", struct_types)
        if not stype:
            break

        width = input_default("가로 (m)", "3")
        depth = input_default("세로 (m)", "3")
        height = input_default("높이 (m)", "2.5")
        spec = f"{width}m × {depth}m × {height}m"

        frame = input_choice("프레임 재질", FRAME_MATERIALS)
        fabric = input_choice("천막 원단", FABRIC_TYPES)
        color = input_choice("색상", COLORS)
        material_desc = f"{frame}/{fabric}/{color}"

        qty = int(input_default("수량", "1"))
        area = float(width) * float(depth)
        base = PRICES["천막구조물"][stype]["base_price"]
        unit_price = int(base * area)
        subtotal = unit_price * qty

        items.append({
            "type": stype, "spec": spec, "material": material_desc,
            "qty": qty, "unit_price": unit_price, "subtotal": subtotal,
        })
        print(f"  → {stype} {spec} × {qty} = {fmt_money(subtotal)}")

        if not input_yn("품목 추가?", "n"):
            break
        idx += 1

    if not items:
        print("  품목이 없습니다. 취소합니다.")
        return

    # 부대비용
    print("\n[ 부대비용 ]")
    install_fee = int(input_default("시공비 (원, 없으면 0)", "0"))
    transport_fee = int(input_default("운반비 (원, 없으면 0)", "0"))
    demolish_fee = int(input_default("철거비 (원, 없으면 0)", "0"))

    # 문서 생성
    quote_number = get_next_quote_number()
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_header(doc, quote_number, "천막구조물")
    add_customer_info(doc, [
        ("고객명", customer),
        ("현장주소", address),
        ("시공일정", schedule),
    ])

    headers = ["No", "구조물종류", "규격", "자재", "수량", "단가", "소계"]
    rows = []
    item_total = 0
    for i, it in enumerate(items, 1):
        rows.append([
            i, it["type"], it["spec"], it["material"],
            it["qty"], fmt_money(it["unit_price"]), fmt_money(it["subtotal"]),
        ])
        item_total += it["subtotal"]

    add_item_table(doc, headers, rows)
    total = add_summary(doc, item_total, install_fee, transport_fee, demolish_fee)
    add_footer(doc)

    # 합계 데이터 구성
    extra = install_fee + transport_fee + demolish_fee
    supply = item_total + extra
    vat = int(supply * 0.1)
    summary_data = [("공급가액 (품목 소계)", fmt_money(item_total))]
    if install_fee:
        summary_data.append(("시공비", fmt_money(install_fee)))
    if transport_fee:
        summary_data.append(("운반비", fmt_money(transport_fee)))
    if demolish_fee:
        summary_data.append(("철거비", fmt_money(demolish_fee)))
    summary_data.append(("부가세 (10%)", fmt_money(vat)))
    summary_data.append(("총 합계", fmt_money(total)))

    print(f"\n  💰 총 합계: {fmt_money(total)}")
    save_document(doc, quote_number, "천막구조물")
    save_excel(quote_number, "천막구조물",
               [("고객명", customer), ("현장주소", address), ("시공일정", schedule)],
               {"headers": headers, "rows": rows}, summary_data)


# ============================================
# 글램핑 견적서
# ============================================
def create_glamping_quote():
    print("\n" + "=" * 50)
    print("  ⛺  글램핑 견적서 작성")
    print("=" * 50)

    # 고객 정보
    print("\n[ 고객 정보 ]")
    customer = input_default("고객명")
    biz_name = input_default("사업장명")
    site_area = input_default("부지면적 (㎡)")

    # 텐트 입력
    items = []
    print("\n[ 텐트 품목 ] (완료시 빈 칸 Enter)")
    idx = 1
    while True:
        print(f"\n  --- 텐트 {idx} ---")
        tent_types = list(PRICES["글램핑"].keys())
        ttype = input_choice("텐트 종류", tent_types)

        model = input_default("모델명", f"{ttype} 기본형")
        spec = input_default("규격", "5m × 7m")
        qty = int(input_default("수량", "1"))

        base = PRICES["글램핑"][ttype]["base_price"]
        subtotal = base * qty

        items.append({
            "type": ttype, "model": model, "spec": spec,
            "qty": qty, "unit_price": base, "subtotal": subtotal,
        })
        print(f"  → {ttype} ({model}) × {qty} = {fmt_money(subtotal)}")

        if not input_yn("텐트 추가?", "n"):
            break
        idx += 1

    if not items:
        print("  품목이 없습니다. 취소합니다.")
        return

    # 옵션
    options = []
    total_qty = sum(it["qty"] for it in items)
    print("\n[ 옵션 선택 ]")
    for opt_name, opt_info in PRICES["글램핑옵션"].items():
        if input_yn(f"{opt_name} ({fmt_money(opt_info['base_price'])}/동) 포함?", "n"):
            opt_total = opt_info["base_price"] * total_qty
            options.append({
                "name": opt_name, "qty": total_qty,
                "unit_price": opt_info["base_price"], "subtotal": opt_total,
            })

    # 문서 생성
    quote_number = get_next_quote_number()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_header(doc, quote_number, "글램핑")
    add_customer_info(doc, [
        ("고객명", customer),
        ("사업장명", biz_name),
        ("부지면적", f"{site_area} ㎡"),
    ])

    # 텐트 테이블
    headers = ["No", "텐트종류", "모델명", "규격", "수량", "단가", "소계"]
    rows = []
    item_total = 0
    for i, it in enumerate(items, 1):
        rows.append([
            i, it["type"], it["model"], it["spec"],
            it["qty"], fmt_money(it["unit_price"]), fmt_money(it["subtotal"]),
        ])
        item_total += it["subtotal"]

    add_item_table(doc, headers, rows, "■ 텐트 내역")

    # 옵션 테이블
    if options:
        opt_headers = ["No", "옵션항목", "수량", "단가", "소계"]
        opt_rows = []
        opt_total = 0
        for i, opt in enumerate(options, 1):
            opt_rows.append([
                i, opt["name"], opt["qty"],
                fmt_money(opt["unit_price"]), fmt_money(opt["subtotal"]),
            ])
            opt_total += opt["subtotal"]
        add_item_table(doc, opt_headers, opt_rows, "■ 옵션 내역")
        item_total += opt_total

    total = add_summary(doc, item_total, 0, 0, 0)
    add_footer(doc)

    vat = int(item_total * 0.1)
    summary_data = [
        ("공급가액 (품목 소계)", fmt_money(item_total)),
        ("부가세 (10%)", fmt_money(vat)),
        ("총 합계", fmt_money(total)),
    ]

    # 글램핑 엑셀용 행 (옵션 포함)
    all_rows = list(rows)
    if options:
        for i, opt in enumerate(options, len(rows) + 1):
            all_rows.append([i, opt["name"], "-", "-", opt["qty"],
                             fmt_money(opt["unit_price"]), fmt_money(opt["subtotal"])])

    print(f"\n  💰 총 합계: {fmt_money(total)}")
    save_document(doc, quote_number, "글램핑")
    save_excel(quote_number, "글램핑",
               [("고객명", customer), ("사업장명", biz_name), ("부지면적", f"{site_area} ㎡")],
               {"headers": headers, "rows": all_rows}, summary_data)


# ============================================
# CLI 원라인 실행 (자동화)
# ============================================
def quick_tent_quote(customer, phone="", address="", date="협의 후 결정"):
    """원라인 천막구조물 견적서 — 기본 어닝 3×3 1개"""
    quote_number = get_next_quote_number()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    customer_info = [("고객명", customer), ("연락처", phone), ("현장주소", address), ("시공일정", date)]
    add_header(doc, quote_number, "천막구조물")
    add_customer_info(doc, customer_info)

    unit_price = PRICES["천막구조물"]["어닝"]["base_price"] * 9  # 3×3
    headers = ["No", "구조물종류", "규격", "자재", "수량", "단가", "소계"]
    rows = [[1, "어닝", "3m × 3m × 2.5m", "알루미늄/PVC코팅/화이트", 1, fmt_money(unit_price), fmt_money(unit_price)]]
    add_item_table(doc, headers, rows)
    total = add_summary(doc, unit_price, 0, 0, 0)
    add_footer(doc)

    vat = int(unit_price * 0.1)
    summary_data = [
        ("공급가액 (품목 소계)", fmt_money(unit_price)),
        ("부가세 (10%)", fmt_money(vat)),
        ("총 합계", fmt_money(total)),
    ]
    save_document(doc, quote_number, "천막구조물")
    save_excel(quote_number, "천막구조물", customer_info,
               {"headers": headers, "rows": rows}, summary_data)
    return quote_number


def quick_glamping_quote(customer, phone="", address="", date="협의 후 결정"):
    """원라인 글램핑 견적서 — 기본 사파리텐트 1동"""
    quote_number = get_next_quote_number()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    customer_info = [("고객명", customer), ("연락처", phone), ("사업장명", address), ("시공일정", date)]
    add_header(doc, quote_number, "글램핑")
    add_customer_info(doc, customer_info)

    unit_price = PRICES["글램핑"]["사파리텐트"]["base_price"]
    headers = ["No", "텐트종류", "모델명", "규격", "수량", "단가", "소계"]
    rows = [[1, "사파리텐트", "사파리텐트 기본형", "5m × 7m", 1, fmt_money(unit_price), fmt_money(unit_price)]]
    add_item_table(doc, headers, rows)
    total = add_summary(doc, unit_price, 0, 0, 0)
    add_footer(doc)

    vat = int(unit_price * 0.1)
    summary_data = [
        ("공급가액 (품목 소계)", fmt_money(unit_price)),
        ("부가세 (10%)", fmt_money(vat)),
        ("총 합계", fmt_money(total)),
    ]
    save_document(doc, quote_number, "글램핑")
    save_excel(quote_number, "글램핑", customer_info,
               {"headers": headers, "rows": rows}, summary_data)
    return quote_number


def list_quotes():
    """생성된 견적서 목록"""
    OUTPUT_DIR.mkdir(exist_ok=True)
    files = sorted(OUTPUT_DIR.glob("WOCS-*"))
    if not files:
        print("  생성된 견적서가 없습니다.")
        return

    print(f"\n  📁 견적서 목록 ({OUTPUT_DIR})")
    print("  " + "-" * 60)
    seen = set()
    for f in files:
        stem = f.stem  # WOCS-2026-001_천막구조물
        if stem not in seen:
            seen.add(stem)
            exts = [x.suffix for x in OUTPUT_DIR.glob(f"{stem}.*")]
            size_kb = sum(x.stat().st_size for x in OUTPUT_DIR.glob(f"{stem}.*")) // 1024
            print(f"  {stem}  ({', '.join(exts)})  {size_kb}KB")
    print(f"\n  총 {len(seen)}건")


# ============================================
# 메인
# ============================================
def main():
    import argparse
    parser = argparse.ArgumentParser(description="WOCS 견적서 자동 생성기")
    parser.add_argument("--type", choices=["awning", "glamping", "both"],
                        help="견적서 종류 (awning=천막구조물, glamping=글램핑, both=둘다)")
    parser.add_argument("--customer", help="고객명")
    parser.add_argument("--phone", default="", help="연락처")
    parser.add_argument("--address", default="", help="현장주소/사업장명")
    parser.add_argument("--date", default="협의 후 결정", help="시공일정")
    parser.add_argument("--list", action="store_true", help="생성된 견적서 목록")
    args = parser.parse_args()

    # --list
    if args.list:
        list_quotes()
        return

    # 원라인 실행
    if args.type and args.customer:
        print()
        print("╔══════════════════════════════════════════╗")
        print("║     WOCS 견적서 자동 생성기 v1.0        ║")
        print("╚══════════════════════════════════════════╝")

        if args.type in ("awning", "both"):
            qn = quick_tent_quote(args.customer, args.phone, args.address, args.date)
            print(f"\n  천막구조물 견적 완료: {qn}")
        if args.type in ("glamping", "both"):
            qn = quick_glamping_quote(args.customer, args.phone, args.address, args.date)
            print(f"\n  글램핑 견적 완료: {qn}")

        print("\n  견적서 생성 완료! quotes/ 폴더를 확인하세요.\n")
        return

    # 대화형 실행
    print()
    print("╔══════════════════════════════════════════╗")
    print("║     WOCS 견적서 자동 생성기 v1.0        ║")
    print("║     우성어닝천막공사캠프시스템            ║")
    print("╚══════════════════════════════════════════╝")

    quote_type = input_choice("견적서 종류", ["천막구조물", "글램핑"])

    if quote_type == "천막구조물":
        create_tent_quote()
    else:
        create_glamping_quote()

    print("\n  견적서 생성 완료! quotes/ 폴더를 확인하세요.\n")


if __name__ == "__main__":
    main()
